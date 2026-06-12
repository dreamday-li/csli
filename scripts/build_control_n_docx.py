from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SRC = Path('artifacts/control_n_report_json_mapping_final.doc')
OUT = Path('artifacts/控N分析报告_JSON结构映射说明_最终版.docx')


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_run_font(run, east_asia='宋体', size=10.5, bold=None, color=None):
    run.font.name = east_asia
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_text_with_inline(paragraph, node):
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = paragraph.add_run(text)
            set_run_font(run)
        return
    if not isinstance(node, Tag):
        return
    if node.name == 'br':
        paragraph.add_run().add_break()
        return
    if node.name in {'strong', 'b'}:
        run = paragraph.add_run(node.get_text())
        set_run_font(run, bold=True)
        return
    if node.name in {'code', 'span'} and ('path' in node.get('class', []) or node.name == 'code'):
        run = paragraph.add_run(node.get_text())
        set_run_font(run, east_asia='Consolas', size=9.5, color=(156, 0, 6))
        return
    for child in node.children:
        add_text_with_inline(paragraph, child)


def add_paragraph(doc, tag, style=None):
    p = doc.add_paragraph(style=style)
    text = tag.get_text(' ', strip=True)
    if not text:
        return p
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for child in tag.children:
        add_text_with_inline(p, child)
    return p


def add_table(doc, table_tag):
    rows = table_tag.find_all('tr', recursive=False)
    if not rows:
        rows = table_tag.find_all('tr')
    max_cols = max((len(r.find_all(['th', 'td'], recursive=False)) for r in rows), default=1)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    table.autofit = True
    for i, row_tag in enumerate(rows):
        cells = row_tag.find_all(['th', 'td'], recursive=False)
        for j, cell_tag in enumerate(cells):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if cell_tag.name == 'th' else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(cell_tag.get_text(' ', strip=True))
            set_run_font(run, size=9.0, bold=(cell_tag.name == 'th'))
            if cell_tag.name == 'th':
                set_cell_shading(cell, 'D9EAF7')
    doc.add_paragraph()


def apply_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = '宋体'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    normal.font.size = Pt(10.5)

    for name, size, color in [
        ('Title', 22, (31, 78, 121)),
        ('Heading 1', 16, (31, 78, 121)),
        ('Heading 2', 13, (47, 85, 151)),
        ('Heading 3', 11.5, (54, 95, 145)),
    ]:
        style = doc.styles[name]
        style.font.name = '黑体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(*color)


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run('控N分析报告 JSON结构映射说明')
        set_run_font(run, size=9, color=(100, 100, 100))
        footer = section.footer
        p2 = footer.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run('报告章节与JSON字段映射说明')
        set_run_font(run2, size=8.5, color=(130, 130, 130))


def main():
    html = SRC.read_text(encoding='utf-8')
    soup = BeautifulSoup(html, 'html.parser')
    doc = Document()
    apply_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    body = soup.body or soup
    for node in body.children:
        if isinstance(node, NavigableString):
            continue
        if not isinstance(node, Tag):
            continue
        classes = node.get('class', [])
        if node.name == 'h1':
            p = doc.add_paragraph(style='Title')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(node.get_text(' ', strip=True))
            set_run_font(run, east_asia='黑体', size=22, bold=True, color=(31,78,121))
        elif node.name == 'h2':
            add_paragraph(doc, node, style='Heading 1')
        elif node.name == 'h3':
            add_paragraph(doc, node, style='Heading 2')
        elif node.name == 'h4':
            add_paragraph(doc, node, style='Heading 3')
        elif node.name == 'p':
            p = add_paragraph(doc, node)
            if 'cover-sub' in classes or 'cover-note' in classes:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif node.name == 'div' and 'pagebreak' in classes:
            doc.add_page_break()
        elif node.name == 'div' and 'box' in classes:
            p = add_paragraph(doc, node)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.right_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        elif node.name == 'div' and 'toc' in classes:
            for child in node.find_all('div', recursive=False):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.2)
                run = p.add_run(child.get_text(' ', strip=True))
                set_run_font(run, size=11)
        elif node.name == 'table':
            add_table(doc, node)
        elif node.name == 'pre':
            p = doc.add_paragraph()
            run = p.add_run(node.get_text())
            set_run_font(run, east_asia='Consolas', size=9)
        elif node.name in {'ul', 'ol'}:
            style = 'List Bullet' if node.name == 'ul' else 'List Number'
            for li in node.find_all('li', recursive=False):
                p = doc.add_paragraph(style=style)
                run = p.add_run(li.get_text(' ', strip=True))
                set_run_font(run)

    add_header_footer(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f'Generated {OUT}')


if __name__ == '__main__':
    main()
